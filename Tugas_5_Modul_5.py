import os
from docx import Document

# Nama file untuk menyimpan data
DATA_FILE = "pesanan.docx"

# Membuat dokumen jika belum ada1

if not os.path.exists(DATA_FILE):
    doc = Document()
    doc.add_heading("Data Pesanan Restoran", level=1)
    doc.save(DATA_FILE)

# Fungsi untuk menambahkan pesanan
def tambah_pesanan():
    nama = input("Masukkan nama pelanggan: ")
    menu = input("Masukkan menu yang dipesan: ")
    jumlah = input("Masukkan jumlah pesanan: ")
    gambar = input("Masukkan path gambar (atau tekan Enter untuk skip): ").strip()

    # Membuka dokumen Word
    doc = Document(DATA_FILE)
    
    # Menambahkan informasi pesanan
    doc.add_paragraph(f"Nama: {nama}, Menu: {menu}, Jumlah: {jumlah}, Status: Diproses")
    
    # Jika path gambar tidak kosong, tambahkan gambar
    if gambar:
        if not os.path.exists(gambar):
            print("Path gambar tidak ditemukan! Pesanan tetap disimpan tanpa gambar.")
        else:
            try:
                doc.add_picture(gambar)  # Menambahkan gambar jika ada
            except Exception as e:
                print(f"Terjadi kesalahan saat menambahkan gambar: {e}. Pesanan tetap disimpan tanpa gambar.")
    
    # Menyimpan dokumen
    doc.save(DATA_FILE)
    print("Pesanan berhasil ditambahkan dan disimpan.")

# Fungsi untuk menampilkan semua pesanan
def tampilkan_pesanan():
    doc = Document(DATA_FILE)
    print("\nData Pesanan:")
    for para in doc.paragraphs[1:]:
        print(para.text)

# Fungsi untuk memperbarui status pesanan
def update_pesanan():
    nama = input("Masukkan nama pelanggan yang ingin diupdate: ")
    doc = Document(DATA_FILE)
    for para in doc.paragraphs:
        if nama in para.text:
            status_baru = input("Masukkan status baru (Proses/Selesai/Batal): ")
            data_baru = para.text.replace("Diproses", status_baru)
            para.text = data_baru
            doc.save(DATA_FILE)
            print("Pesanan berhasil diperbarui.")
            return
    print("Pesanan tidak ditemukan.")

# Fungsi untuk menghapus pesanan dengan status 'Batal'
def hapus_pesanan():
    nama = input("Masukkan nama pelanggan yang ingin dihapus: ")

    doc = Document(DATA_FILE)
    paragraphs_to_keep = []  # Daftar untuk menyimpan paragraf yang tidak dihapus
    deleted = False

    for para in doc.paragraphs:
        # Periksa apakah nama pelanggan ada dan statusnya 'Batal'
        if nama in para.text and "Batal" in para.text:
            deleted = True  # Tandai bahwa pesanan telah dihapus
        else:
            paragraphs_to_keep.append(para.text)

    if deleted:
        # Hapus semua paragraf dari dokumen
        for _ in doc.paragraphs:
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)

        # Tambahkan kembali paragraf yang tetap
        for text in paragraphs_to_keep:
            doc.add_paragraph(text)

        doc.save(DATA_FILE)
        print("Pesanan berhasil dihapus.")
    else:
        print("Pesanan tidak ditemukan atau status tidak Batal.")

# Fungsi untuk mencari pesanan berdasarkan nama pelanggan
def cari_pesanan():
    nama = input("Masukkan nama pelanggan yang ingin dicari: ")
    doc = Document(DATA_FILE)
    for para in doc.paragraphs:
        if nama in para.text:
            print(f"Ditemukan: {para.text}")
            return
    print("Pesanan tidak ditemukan.")

# Menu utama
def menu():
    while True:
        print("\nManajemen Pesanan Restoran")
        print("1. Tambah Pesanan")
        print("2. Tampilkan Pesanan")
        print("3. Update Pesanan")
        print("4. Hapus Pesanan")
        print("5. Cari Pesanan (Opsional)")
        print("6. Keluar")

        pilihan = input("Pilih menu: ")
        if pilihan == "1":
            tambah_pesanan()
        elif pilihan == "2":
            tampilkan_pesanan()
        elif pilihan == "3":
            update_pesanan()
        elif pilihan == "4":
            hapus_pesanan()
        elif pilihan == "5":
            cari_pesanan()
        elif pilihan == "6":
            print("Terima kasih!")
            break
        else:
            print("Pilihan tidak valid. Coba lagi.")

# Menjalankan program
menu()
