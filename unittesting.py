import os
import unittest
from docx import Document
from Tugas_5_Modul_5 import tambah_pesanan, tampilkan_pesanan, update_pesanan, hapus_pesanan, cari_pesanan

TEST_FILE = "test_pesanan.docx"

# Membuat setup dan teardown untuk pengujian
class TestManajemenPesanan(unittest.TestCase):
    def setUp(self):
        """Dijalankan sebelum setiap tes, membuat file dokumen sementara."""
        self.test_file = TEST_FILE
        if os.path.exists(self.test_file):
            os.remove(self.test_file)
        doc = Document()
        doc.add_heading("Data Pesanan Restoran", level=1)
        doc.save(self.test_file)

    def tearDown(self):
        """Dijalankan setelah setiap tes, menghapus file dokumen sementara."""
        if os.path.exists(self.test_file):
            os.remove(self.test_file)

    def test_tambah_pesanan(self):
        """Menguji fungsi tambah_pesanan."""
        input_values = ["John Doe", "Nasi Goreng", "2", ""]
        expected_output = "Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Diproses"
        
        # Mock input
        def mock_input(prompt):
            return input_values.pop(0)

        original_input = __builtins__.input
        __builtins__.input = mock_input

        try:
            tambah_pesanan(self.test_file)
            doc = Document(self.test_file)
            self.assertIn(expected_output, [p.text for p in doc.paragraphs])
        finally:
            __builtins__.input = original_input

    def test_tampilkan_pesanan(self):
        """Menguji fungsi tampilkan_pesanan."""
        doc = Document(self.test_file)
        doc.add_paragraph("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Diproses")
        doc.save(self.test_file)

        expected_output = "Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Diproses"

        # Capture output
        from io import StringIO
        import sys

        captured_output = StringIO()
        sys.stdout = captured_output

        tampilkan_pesanan(self.test_file)
        sys.stdout = sys.__stdout__

        self.assertIn(expected_output, captured_output.getvalue())

    def test_update_pesanan(self):
        """Menguji fungsi update_pesanan."""
        doc = Document(self.test_file)
        doc.add_paragraph("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Diproses")
        doc.save(self.test_file)

        input_values = ["John Doe", "Selesai"]
        def mock_input(prompt):
            return input_values.pop(0)

        original_input = __builtins__.input
        __builtins__.input = mock_input

        try:
            update_pesanan(self.test_file)
            doc = Document(self.test_file)
            self.assertIn("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Selesai", [p.text for p in doc.paragraphs])
        finally:
            __builtins__.input = original_input

    def test_hapus_pesanan(self):
        """Menguji fungsi hapus_pesanan."""
        doc = Document(self.test_file)
        doc.add_paragraph("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Batal")
        doc.add_paragraph("Nama: Jane Doe, Menu: Soto Ayam, Jumlah: 1, Status: Diproses")
        doc.save(self.test_file)

        input_values = ["John Doe"]
        def mock_input(prompt):
            return input_values.pop(0)

        original_input = __builtins__.input
        __builtins__.input = mock_input

        try:
            hapus_pesanan(self.test_file)
            doc = Document(self.test_file)
            paragraphs = [p.text for p in doc.paragraphs]
            self.assertNotIn("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Batal", paragraphs)
            self.assertIn("Nama: Jane Doe, Menu: Soto Ayam, Jumlah: 1, Status: Diproses", paragraphs)
        finally:
            __builtins__.input = original_input

    def test_cari_pesanan(self):
        """Menguji fungsi cari_pesanan."""
        doc = Document(self.test_file)
        doc.add_paragraph("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Diproses")
        doc.save(self.test_file)

        input_values = ["John Doe"]
        def mock_input(prompt):
            return input_values.pop(0)

        original_input = __builtins__.input
        __builtins__.input = mock_input

        try:
            from io import StringIO
            import sys

            captured_output = StringIO()
            sys.stdout = captured_output

            cari_pesanan(self.test_file)
            sys.stdout = sys.__stdout__

            self.assertIn("Nama: John Doe, Menu: Nasi Goreng, Jumlah: 2, Status: Diproses", captured_output.getvalue())
        finally:
            __builtins__.input = original_input

# Menjalankan tes
if __name__ == "__main__":
    unittest.main()
